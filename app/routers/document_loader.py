
import hashlib  # Импортируем стандартную библиотеку для генерации хэшей
import os       # Импортируем для работы с путями к файлам
import io       # Импортируем для работы с бинарными потоками файлов

from dependencies import chroma, record_manager, database, xml_documents_table # Импортируем database и xml_documents_table
from schemas.document_upload import DocumentUpload # Схема для загрузки документа в виде JSON
from schemas.db_models import XMLDocument # Импортируем Pydantic-схему для XMLDocument

from fastapi import APIRouter, Body, Query, UploadFile, File, HTTPException, status # Добавляем UploadFile, File, HTTPException, status

from langchain_core.indexing import aindex
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# Импорты для обработки файлов
from langchain_community.document_loaders import PyPDFLoader # Для PDF
import docx # Для DOCX


router = APIRouter()

# Инициализация разделителя текста на чанки
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1500,
    chunk_overlap=150,
    length_function=len,
    is_separator_regex=False,
    separators=["\n", "\n\n", "."],
)

# Эндпоинт для загрузки документа в виде JSON-объекта
@router.post(
    "/load/doc_as_json",
    description="Загрузка документа - в виде JSON-объекта - с детерминированным ID для защиты от дубликатов"
)
async def upload_document(document_input: DocumentUpload = Body()):
    # 1. Генерируем стабильный parent_doc_id на основе хэша текста (SHA-256)
    # Для одного и того же текста этот ID всегда будет одинаковым
    text_hash = hashlib.sha256(document_input.content.encode("utf-8")).hexdigest()
    parent_id = text_hash

    # 2. Формируем метаданные, которые будут уникальны для этого контента
    final_metadata = document_input.metadata.copy() if document_input.metadata else {}
    final_metadata["parent_doc_id"] = parent_id

    # 3. Нарезаем текст на чанки
    chunks: list[Document] = text_splitter.create_documents(
        texts=[document_input.content],
        metadatas=[final_metadata]
    )

    # 4. Передаем данные в aindex. Теперь при повторном запросе хэши чанков сойдутся,
    # и record_manager заблокирует дублирование в Chroma
    out = await aindex(
        chunks,
        record_manager=record_manager,
        vector_store=chroma
    )

    return {
        "index_status": out,
        "parent_doc_id": parent_id
    }

# Эндпоинт для получения списка всех чанков
@router.get(
    "/get/documents",
    description="Получить список всех чанков и их метаданных, записанных в Chroma"
)
async def get_stored_documents(
    parent_doc_id: str | None = Query(default=None, description="Фильтр по ID родительского документа"),
    limit: int = Query(default=10, description="Количество выводимых чанков")
):
    try:
        # 1. Формируем фильтр, если клиент передал parent_doc_id
        where_filter = {"parent_doc_id": parent_doc_id} if parent_doc_id else None

        # 2. Делаем прямой запрос к коллекции Chroma через метод .get()
        # Извлекаем тексты (documents) и метаданные (metadatas)
        result = chroma.get(
            where=where_filter,
            limit=limit,
            include=["documents", "metadatas"]  # Можно добавить "embeddings", если нужны сами векторы-числа
        )

        # 3. Форматируем ответ для удобного чтения на фронтенде или в Swagger
        formatted_chunks = []
        if result and "documents" in result:
            for idx in range(len(result["ids"])):
                formatted_chunks.append({
                    "chunk_id": result["ids"][idx],
                    "text": result["documents"][idx],
                    "metadata": result["metadatas"][idx] if result["metadatas"] else {}
                })

        return {
            "total_fetched": len(formatted_chunks),
            "chunks": formatted_chunks
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка чтения из Chroma: {str(e)}")


# Вспомогательная функция для чтения содержимого файла в зависимости от его типа
async def _read_document_content(file: UploadFile) -> str:
    """Вспомогательная функция для чтения содержимого файла в зависимости от его типа.
    Поддерживает .txt, .pdf, .docx."""
    content = ""
    # Определяем расширение файла по его имени
    file_extension = os.path.splitext(file.filename)[1].lower()

    if file_extension == ".txt":
        # Для текстовых файлов просто читаем содержимое
        content = (await file.read()).decode("utf-8")
    elif file_extension == ".pdf":
        # Для PDF используем PyPDFLoader из LangChain
        # PyPDFLoader требует путь к файлу, поэтому сохраним его временно
        temp_pdf_path = f"/tmp/{file.filename}"
        with open(temp_pdf_path, "wb") as buffer:
            buffer.write(await file.read())
        loader = PyPDFLoader(temp_pdf_path)
        pages = loader.load()
        # Объединяем содержимое всех страниц
        content = "\n".join([p.page_content for p in pages])
        os.remove(temp_pdf_path) # Удаляем временный файл
    elif file_extension == ".docx":
        # Для DOCX используем python-docx, читая из байтового потока
        doc = docx.Document(io.BytesIO(await file.read()))
        # Извлекаем текст из каждого параграфа
        content = "\n".join([para.text for para in doc.paragraphs])
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Неподдерживаемый формат файла: {file_extension}. Поддерживаются .txt, .pdf, .docx."
        )
    return content


@router.post(
    "/load/doc_as_file_and_xml",
    description="""Загрузка основного документа (TXT, PDF, DOCX) и связанного с ним XML файла.
                 Основной документ индексируется в Chroma, XML сохраняется в PostgreSQL."""
)
async def upload_file_and_xml(
    document_file: UploadFile = File(..., description="Основной документ (TXT, PDF, DOCX)"),
    xml_file: UploadFile = File(..., description="Связанный XML-файл")
):
    try:
        # 1. Читаем содержимое основного документа
        main_doc_content = await _read_document_content(document_file)
        if not main_doc_content:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Основной документ пуст или не удалось прочитать."
            )

        # 2. Генерируем стабильный parent_doc_id на основе хэша основного документа
        parent_id = hashlib.sha256(main_doc_content.encode("utf-8")).hexdigest()

        # 3. Читаем содержимое XML файла
        xml_content = (await xml_file.read()).decode("utf-8")
        if not xml_content:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="XML-файл пуст или не удалось прочитать."
            )

        # 4. Сохраняем XML документ в PostgreSQL
        xml_doc_data = XMLDocument(
            parent_doc_id=parent_id,
            filename=xml_file.filename,
            xml_content=xml_content
        )
        query = xml_documents_table.insert().values(**xml_doc_data.model_dump())
        await database.execute(query)
        print(f"XML-документ '{xml_file.filename}' сохранен в PostgreSQL с parent_doc_id: {parent_id}")

        # 5. Формируем метаданные для чанков основного документа
        final_metadata = {
            "parent_doc_id": parent_id,
            "original_filename": document_file.filename,
            "xml_filename": xml_file.filename
        }

        # 6. Нарезаем основной документ на чанки и индексируем их в Chroma
        chunks: list[Document] = text_splitter.create_documents(
            texts=[main_doc_content],
            metadatas=[final_metadata]
        )

        out = await aindex(
            chunks,
            record_manager=record_manager,
            vector_store=chroma
        )
        print(f"Основной документ '{document_file.filename}' проиндексирован в Chroma с parent_doc_id: {parent_id}")

        return {
            "index_status": out,
            "parent_doc_id": parent_id,
            "message": "Документ и XML успешно загружены и проиндексированы."
        }

    except HTTPException as e:
        # Перехватываем и повторно выбрасываем HTTPException для корректной обработки FastAPI
        raise e
    except Exception as e:
        # Обрабатываем другие непредвиденные ошибки
        print(f"Ошибка при загрузке файла и XML: {e}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Внутренняя ошибка сервера: {str(e)}")
